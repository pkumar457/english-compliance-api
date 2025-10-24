
import io
from typing import Tuple
from docx import Document as DocxDocument
from pypdf import PdfReader
from pdfminer.high_level import extract_text as pdfminer_extract

def _text_from_docx_bytes(b: bytes) -> str:
    f = io.BytesIO(b); doc = DocxDocument(f)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.append(" ".join(c.text for c in row.cells if c.text))
    return "\n".join([p for p in parts if p and p.strip()])

def _text_from_pdf_bytes(b: bytes) -> str:
    try:
        f = io.BytesIO(b); reader = PdfReader(f)
        parts = [(p.extract_text() or "") for p in reader.pages]
        text = "\n".join(parts).strip()
        if len(text) > 20: return text
    except Exception: pass
    return pdfminer_extract(io.BytesIO(b))

def extract_text(filename: str, file_bytes: bytes) -> Tuple[str, str]:
    lower = filename.lower()
    if lower.endswith(".docx"): return "docx", _text_from_docx_bytes(file_bytes)
    if lower.endswith(".pdf"): return "pdf", _text_from_pdf_bytes(file_bytes)
    raise ValueError("Unsupported file type. Please upload .pdf or .docx")
